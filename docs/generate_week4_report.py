#!/usr/bin/env python3
"""Generate Week 4 implementation summary as Word (.docx)."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

OUT = Path(__file__).resolve().parent / "Week-4-Implementation-Summary.docx"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val


def build() -> None:
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("UniSchedule AI\nWeek 4 — Implementation Summary")
    run.bold = True
    run.font.size = Pt(18)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(
        "Intelligent University Meeting Scheduler\n"
        "Final Year Project · University of Canterbury\n"
        "May 2026"
    )

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.add_run("Student: ").bold = True
    meta.add_run("[Your Full Name]\n")
    meta.add_run("Student ID: ").bold = True
    meta.add_run("[Student ID]\n")
    meta.add_run("Supervisor: ").bold = True
    meta.add_run("[Supervisor Name]")

    doc.add_page_break()

    add_heading(doc, "1. Purpose of this document", 1)
    doc.add_paragraph(
        "This document summarises work completed in Week 4 of the UniSchedule AI "
        "project. It describes the implemented software artefacts (backend API, "
        "database, frontend shell, and clash-detection logic), how they align with "
        "the Week 4 development plan, and what remains for later weeks. It is "
        "intended for supervisor review and dissertation progress evidence."
    )

    add_heading(doc, "2. Week 4 objectives (from project plan)", 1)
    doc.add_paragraph("Week 4 targeted the following outcomes:")
    add_bullets(
        doc,
        [
            "Database migrations for core scheduling entities (MySQL / SQLite).",
            "Laravel REST API with Laravel Sanctum token authentication and student/lecturer roles.",
            "Scheduling service with meeting create, read, update, and cancel (CRUD).",
            "Real-time clash detection across participants, availability, rooms, and policy rules.",
            "Next.js frontend shell: login, register, dashboard, and meeting booking form.",
            "Research: ethics pathway, survey launch, and interview scheduling (parallel track).",
        ],
    )

    add_heading(doc, "3. Technology stack (implemented)", 1)
    add_table(
        doc,
        ["Layer", "Technology", "Role in Week 4"],
        [
            ["Frontend", "Next.js 16 + TypeScript", "User interface; calls Laravel API"],
            ["Backend", "Laravel 11 (PHP)", "REST API, business logic, validation"],
            ["Database", "MySQL or SQLite", "Persistent storage via Eloquent ORM"],
            ["Authentication", "Laravel Sanctum", "Bearer API tokens (not cookie SPA mode)"],
            ["Architecture", "Decoupled web app", "Next.js → HTTPS → Laravel API → database"],
        ],
    )

    add_heading(doc, "4. Database schema implemented", 1)
    doc.add_paragraph(
        "Migrations create the following tables with relationships suitable for "
        "academic-aware scheduling and evaluation:"
    )
    add_table(
        doc,
        ["Table", "Purpose"],
        [
            ["users", "Accounts with role: student or lecturer"],
            ["rooms", "Bookable spaces (name, building, capacity)"],
            ["availabilities", "Per-user weekly availability windows (day + time range)"],
            ["constraint_rules", "Academic policies (e.g. exam blackout periods)"],
            ["meetings", "Scheduled meetings (organiser, room, times, status)"],
            ["meeting_participants", "Users linked to each meeting"],
            ["clash_records", "Logged clash events when booking is blocked or forced"],
            ["personal_access_tokens", "Sanctum API tokens"],
        ],
    )

    add_heading(doc, "5. Backend API — implemented endpoints", 1)
    doc.add_paragraph("Base URL: http://localhost:8000/api/v1")
    add_table(
        doc,
        ["Method", "Endpoint", "Description", "Auth"],
        [
            ["POST", "/auth/register", "Register student or lecturer", "Public"],
            ["POST", "/auth/login", "Login; returns Bearer token", "Public"],
            ["POST", "/auth/logout", "Revoke current token", "Sanctum"],
            ["GET", "/auth/me", "Current user profile", "Sanctum"],
            ["GET", "/meetings", "List meetings for current user", "Sanctum"],
            ["POST", "/meetings", "Create meeting (clash-checked)", "Sanctum"],
            ["GET", "/meetings/{id}", "View single meeting", "Sanctum"],
            ["PUT", "/meetings/{id}", "Update meeting (organiser only)", "Sanctum"],
            ["DELETE", "/meetings/{id}", "Cancel meeting (organiser only)", "Sanctum"],
            ["POST", "/meetings/check-clash", "Preview clashes + alternative slots", "Sanctum"],
        ],
    )

    add_heading(doc, "6. Core backend services", 1)

    add_heading(doc, "6.1 SchedulingService", 2)
    add_bullets(
        doc,
        [
            "Creates and updates meetings inside a database transaction.",
            "Validates room capacity against participant count.",
            "Runs clash detection before persisting; returns HTTP 422 with clash details if blocked.",
            "Supports optional force flag for research/logging scenarios.",
            "Suggests up to five alternative time slots when clashes are detected.",
        ],
    )

    add_heading(doc, "6.2 ClashDetectionService", 2)
    doc.add_paragraph("Detects the following clash types:")
    add_table(
        doc,
        ["Type", "Rule"],
        [
            ["time", "End time must be after start time"],
            ["participant", "User already has an overlapping scheduled meeting"],
            ["availability", "Meeting outside declared weekly availability hours"],
            ["room", "Room double-booked for overlapping interval"],
            ["policy", "Active constraint rules (e.g. exam blackout) block the date"],
        ],
    )

    add_heading(doc, "6.3 Validation and security", 2)
    add_bullets(
        doc,
        [
            "Form Request classes validate auth and meeting input.",
            "API Resources format consistent JSON responses.",
            "Role middleware (EnsureUserRole) registered for future route protection.",
            "CORS configured for Next.js origin (localhost:3000).",
            "Bearer token authentication; CSRF not required for token-based API clients.",
        ],
    )

    add_heading(doc, "7. Frontend — implemented pages", 1)
    add_table(
        doc,
        ["Route", "Feature"],
        [
            ["/", "Landing page with links to login/register"],
            ["/login", "Email/password login; stores Sanctum token"],
            ["/register", "Registration with role selection (student/lecturer)"],
            ["/dashboard", "Lists user meetings with status and room"],
            ["/dashboard/meetings/new", "Create meeting form + Check clashes button"],
        ],
    )
    doc.add_paragraph(
        "The API client (frontend/src/lib/api.ts) attaches Authorization: Bearer {token} "
        "on all protected requests and displays readable clash error messages."
    )

    add_heading(doc, "8. Demo seed data", 1)
    doc.add_paragraph("Running php artisan db:seed creates:")
    add_bullets(
        doc,
        [
            "Lecturer: lecturer@unischedule.test (password: password)",
            "Student: student@unischedule.test (password: password)",
            "Two rooms: ENG-101 (capacity 30), SCI-202 (capacity 20)",
            "Availability windows for both users (all days, 08:00–20:00 after patch seeder)",
            "Sample exam blackout rule (inactive by default for demo)",
        ],
    )

    add_heading(doc, "9. Automated tests", 1)
    add_bullets(
        doc,
        [
            "Feature tests: ClashDetectionTest (participant overlap, non-overlap, availability window).",
            "Run with: php artisan test (from backend folder).",
            "Additional integration and ≥20 evaluation scenarios planned for Week 7.",
        ],
    )

    add_heading(doc, "10. Alignment with MoSCoW (Must-have progress)", 1)
    add_table(
        doc,
        ["Requirement", "Week 4 status"],
        [
            ["Sanctum auth + student/lecturer roles", "Implemented"],
            ["Meeting CRUD", "Implemented"],
            ["Clash detection + academic constraints", "Implemented (v1)"],
            ["Room allocation", "Basic (room_id + capacity check)"],
            ["NLP / hybrid meeting requests", "Not started — planned Week 5"],
            ["Survey + evaluation (manual vs AI)", "Not started — planned Weeks 4–10"],
            ["Alternative slot suggestions", "Implemented on check-clash endpoint"],
        ],
    )

    add_heading(doc, "11. Research and documentation (Week 4 parallel track)", 1)
    doc.add_paragraph(
        "The following research deliverables were planned for Week 4. Update checkboxes "
        "after supervisor meetings:"
    )
    add_table(
        doc,
        ["Item", "Status"],
        [
            ["Week 3 supervisor submission pack (6 PDFs)", "Completed and in docs/supervisor-submission/"],
            ["Literature review Chapter 2 draft", "In progress"],
            ["Ethics application for survey/interviews", "To be confirmed with supervisor"],
            ["Survey instrument live", "Planned — not yet deployed"],
            ["Interviews scheduled (3–5)", "Planned"],
            ["≥20 formal test scenarios documented", "Planned — Week 4 evaluation milestone"],
        ],
    )

    add_heading(doc, "12. Not yet implemented (future weeks)", 1)
    add_bullets(
        doc,
        [
            "NLP / natural-language meeting booking (Week 5).",
            "Full end-to-end AI vertical slice polish (Week 6).",
            "Analytics dashboard (Week 8).",
            "Calendar sync with Google/Outlook (Should-have; may defer).",
            "Comparative evaluation: manual vs AI-assisted scheduling (Week 10).",
            "Pilot user acceptance testing (Week 9).",
        ],
    )

    add_heading(doc, "13. How to run the prototype", 1)
    add_bullets(
        doc,
        [
            "Backend: cd backend → php artisan serve → http://localhost:8000",
            "Frontend: cd frontend → npm run dev → http://localhost:3000",
            "See SETUP.md in the project root for migrations, seeding, and environment variables.",
        ],
    )

    add_heading(doc, "14. Suggested supervisor demo (5–10 minutes)", 1)
    add_bullets(
        doc,
        [
            "Register or log in as student@unischedule.test.",
            "Create a meeting on a weekday between 08:00 and 20:00 with room ID 1 or 2.",
            "Use Check clashes to preview conflict detection.",
            "Attempt a double-booking or out-of-hours slot to show validation messages.",
            "Show meeting list on dashboard and optional API response in browser dev tools.",
        ],
    )

    add_heading(doc, "15. Summary", 1)
    doc.add_paragraph(
        "Week 4 delivered a working research prototype foundation: a Laravel API with "
        "Sanctum authentication, a relational schema for academic scheduling, deterministic "
        "clash detection and scheduling logic, PHPUnit tests for core clash rules, and a "
        "Next.js frontend for authentication and meeting management. This satisfies the "
        "core development goals for Week 4 and provides a platform for NLP integration "
        "and empirical evaluation in subsequent weeks."
    )

    doc.add_paragraph()
    p = doc.add_paragraph("Document generated from the UniSchedule AI repository — May 2026.")
    p.runs[0].italic = True

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
